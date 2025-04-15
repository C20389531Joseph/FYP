from flask import Blueprint, render_template, request, current_app
import openai
from .modules_BERT import grade_BERT_essay, extract_text
from .models import grade_LSTM_essay
import os
from transformers import pipeline, GPT2Tokenizer
from docx import Document


main = Blueprint('main', __name__)

# Load GPT-2 pipeline and tokenizer
generator = pipeline("text-generation", model="gpt2")
tokenizer = GPT2Tokenizer.from_pretrained("gpt2")


def generate_explanation(user_essay, scores):
    print("Hello")
    intro = "The user submitted the following answer:\n\n"
    scores_part = f"The grading model gave the following scores: {scores}\n\n"
    instructions = "Please provide a clear, constructive explanation for the student to help them improve. Highlight what was done well and what could be improved. Feedback:"

    # Encode static parts
    intro_ids = tokenizer.encode(intro, add_special_tokens=False)
    scores_ids = tokenizer.encode(scores_part, add_special_tokens=False)
    instructions_ids = tokenizer.encode(instructions, add_special_tokens=False)

    static_token_count = len(intro_ids) + len(scores_ids) + len(instructions_ids)
    max_tokens = 1024
    max_new_tokens = 300
    available_for_essay = max_tokens - max_new_tokens - static_token_count
    print("available_for_essay")
    print(available_for_essay)
    # Encode full essay
    essay_ids = tokenizer.encode(user_essay, add_special_tokens=False)

    # Split into chunks
    chunk_size = available_for_essay
    chunks = [essay_ids[i:i + chunk_size] for i in range(0, len(essay_ids), chunk_size)][:3]

    explanations = []

    for chunk in chunks:
        final_input_ids = intro_ids + chunk + scores_ids + instructions_ids

        # Truncate from start if necessary to maintain total 1024 tokens
        max_context_length = 1024 - 350  # leave space for max_new_tokens
        if len(final_input_ids) > max_context_length:
            final_input_ids = final_input_ids[-max_context_length:]  # keep the most recent context
        print("final_input_ids")
        print(len(final_input_ids))
        final_prompt = tokenizer.decode(final_input_ids, skip_special_tokens=True)
        print("final_prompt")
        # Generate
        output = generator(final_prompt, max_new_tokens=300, do_sample=True, temperature=0.7)
        generated_text = output[0]['generated_text']

        # Extract feedback
        explanation_start = generated_text.find("Feedback:")
        explanation = generated_text[explanation_start + len("Feedback:"):].strip() if explanation_start != -1 else generated_text.strip()

        explanation = explanation.split("Score:")[0].strip()  # Clean up
        explanations.append(explanation)

    full_explanation = "\n\n".join(explanations)
    return full_explanation


def get_exam_paper_path(year):
    # Map years to filenames
    paper_map = {
        "2023Adv4007": "2023Adv4007.docx",
        "2024Adv4007": "2024Adv4007.docx",
        "SampleQuestions": "SampleQuestions.docx",
    }
    filename = paper_map.get(str(year))
    if not filename:
        raise FileNotFoundError(f"No paper defined for year {year}")
    print(current_app.root_path)
    base_path = os.path.join(current_app.root_path, 'static', 'exams')
    full_path = os.path.join(base_path, filename)
    if not os.path.exists(full_path):
        raise FileNotFoundError(f"File {full_path} does not exist")

    return full_path

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

@main.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'document' not in request.files:
            return render_template('upload.html', result="No file part in the request.")
        file = request.files['document']
        if file.filename == '':
            return render_template('upload.html', result="No file selected.")

        year = request.form.get('year')

        if file and year:
            filepath = os.path.join('uploads', file.filename)
            file.save(filepath)

            text = extract_text(filepath)  # This is the student's essay
            exam_path = get_exam_paper_path(year)

            # Extract text from exam paper
            if exam_path.endswith(".pdf"):
                import fitz  # PyMuPDF
                with fitz.open(exam_path) as doc:
                    exam_text = "\n".join(page.get_text() for page in doc)
            else:  # Assume DOCX
                import docx
                exam_doc = docx.Document(exam_path)
                exam_text = "\n".join(p.text for p in exam_doc.paragraphs)

            # Grade using LSTM and BERT
            resultLSTM = grade_LSTM_essay(text, year)
            resultBERT = grade_BERT_essay(text, year)

            # Format grading results nicely
            result_html = f"""
            <h4>LSTM Grades:</h4> {resultLSTM}<br>
            <h4>BERT Grades:</h4> {resultBERT}
            """

            # # Generate AI feedback explanation
            try:
                 explanation = generate_explanation(text, result_html)
            except openai.error.InvalidRequestError as e:
                 explanation = "AI feedback could not be generated due to quota limits."

            # Return result to template
            return render_template(
                'upload.html',
                result=result_html,
                explanation=explanation,
                year=year,
                essay=text
            )

    return render_template('upload.html')