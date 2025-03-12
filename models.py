'''
import torch
import tensorflow as tf
from transformers import BertTokenizer, BertForSequenceClassification
import docx
import numpy as np

class EssayGradingModel:
    def __init__(self, model_type):
        self.model_type = model_type
        self.models = {
            "LSTM": tf.keras.models.load_model(r"C:\Users\JOSEP\OneDrive\Documents\FYP\DemoCode\essay_grading_model.keras"),
            # "BERT": BertForSequenceClassification.from_pretrained("bert_essay_grader"),
            # "GPT": BertForSequenceClassification.from_pretrained("gpt_essay_grader"),
        }
        # self.tokenizers = {
        #     "BERT": BertTokenizer.from_pretrained("bert_essay_grader"),
        #     "GPT": BertTokenizer.from_pretrained("gpt_essay_grader"),
        # }

    def process_text(self, file):
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])

    def grade_essay(self, file):
        essay_text = self.process_text(file)

        # if self.model_type in ["BERT", "GPT"]:
        #     tokenizer = self.tokenizers[self.model_type]
        #     tokens = tokenizer(essay_text, padding=True, truncation=True, return_tensors="pt")
        #     with torch.no_grad():
        #         output = self.models[self.model_type](**tokens)
        #     scores = output.logits.squeeze().tolist()
        
        # if self.model_type == "LSTM":
        tokenizer = tf.keras.preprocessing.text.Tokenizer()
        sequence = tokenizer.texts_to_sequences([essay_text])
        padded_sequence = tf.keras.preprocessing.sequence.pad_sequences(sequence, maxlen=500)
        scores = self.models["LSTM"].predict(padded_sequence)[0].tolist()

        rubric_categories = ["Grammar", "Coherence", "Content Depth", "Language Clarity"]
        feedback = {category: f"Score: {score:.2f}" for category, score in zip(rubric_categories, scores)}

        return feedback '''

import torch
import tensorflow as tf
from transformers import BertTokenizer
import docx
import numpy as np

class EssayGradingModel:
    def __init__(self, model_type):
        self.model_type = model_type
        self.models = {
            "LSTM": tf.keras.models.load_model(r"C:\Users\JOSEP\OneDrive\Documents\FYP\DemoCode\essay_grading_model.keras"),
            "BERT": tf.keras.models.load_model(r"C:\Users\JOSEP\Documents\bert_essay_grading_model\bert_model.keras"),
            # "GPT": BertForSequenceClassification.from_pretrained("gpt_essay_grader"),
        }
        self.tokenizers = {
            "BERT": BertTokenizer.from_pretrained("bert-base-uncased"),
        }

    def process_text(self, file):
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])

    def grade_essay(self, file):
        essay_text = self.process_text(file)

        if self.model_type == "BERT":
            tokenizer = self.tokenizers["BERT"]
            tokens = tokenizer(essay_text, padding=True, truncation=True, return_tensors="tf")
            scores = self.models["BERT"].predict(tokens.data)[0].tolist()
        
        elif self.model_type == "LSTM":
            tokenizer = tf.keras.preprocessing.text.Tokenizer()
            sequence = tokenizer.texts_to_sequences([essay_text])
            padded_sequence = tf.keras.preprocessing.sequence.pad_sequences(sequence, maxlen=500)
            scores = self.models["LSTM"].predict(padded_sequence)[0].tolist()
        
        else:
            raise ValueError("Unsupported model type. Choose either 'LSTM' or 'BERT'.")

        rubric_categories = ["Grammar", "Coherence", "Content Depth", "Language Clarity"]
        feedback = {category: f"Score: {score:.2f}" for category, score in zip(rubric_categories, scores)}

        return feedback
